# import pytesseract
# from PIL import Image
# from docx import Document
# from io import BytesIO

# # Configura o caminho para o executável do Tesseract se necessário (no Windows)
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# def extract_images_from_docx(docx_path):
#     doc = Document(docx_path)
#     images = []
    
#     # Percorre todos os elementos no documento
#     for rel in doc.part.rels.values():
#         if "image" in rel.target_ref:
#             image_data = rel.target_part.blob
#             images.append(image_data)
    
#     return images

# def extract_text_from_image(image_data):
#     image = Image.open(BytesIO(image_data))
#     text = pytesseract.image_to_string(image)
#     return text

# # Caminho do arquivo .docx
# docx_path = r"E:\Git\mapas-se\arquivos\Uploads\testedoc3.docx"

# # Extrai as imagens do documento
# images = extract_images_from_docx(docx_path)

# # Extrai o texto de cada imagem
# for idx, image_data in enumerate(images):
#     print(f"Texto extraído da imagem {idx + 1}:")
#     text = extract_text_from_image(image_data)
#     print(text)
#     print("-" * 40)



import pytesseract
from PIL import Image
from docx import Document
from io import BytesIO

# Configura o caminho para o executável do Tesseract se necessário (no Windows)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    all_text = ""
    
    # Extrai o texto padrão do documento
    for para in doc.paragraphs:
        all_text += para.text + "\n"
    
    return all_text

def extract_images_from_docx(docx_path):
    doc = Document(docx_path)
    images = []
    
    # Percorre todos os elementos no documento para encontrar imagens
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            images.append(image_data)
    
    return images

def extract_text_from_image(image_data):
    image = Image.open(BytesIO(image_data))
    text = pytesseract.image_to_string(image)
    return text

def save_text_to_file(text, output_path):
    with open(output_path, 'w', encoding='utf-8') as file:
        file.write(text)

# Caminho do arquivo .docx
docx_path = r"E:\Git\mapas-se\arquivos\Uploads\testedoc3.docx"# Caminho do arquivo de saída .txt
output_path = 'resultadodocx.txt'

# Extrai o texto do documento
docx_text = extract_text_from_docx(docx_path)

# Extrai as imagens do documento
images = extract_images_from_docx(docx_path)

# Extrai o texto de cada imagem e adiciona ao texto geral
for idx, image_data in enumerate(images):
    image_text = extract_text_from_image(image_data)
    docx_text += f"\nTexto extraído da imagem {idx + 1}:\n{image_text}\n"
    docx_text += "-" * 40 + "\n"

# Salva o texto extraído (incluindo o das imagens) em um arquivo .txt
save_text_to_file(docx_text, output_path)

print(f"O texto extraído foi salvo em {output_path}")
